"""Lightweight RAG ingestion and retrieval for local literature.

Reads documents from a folder, chunks them, builds embeddings, and returns the
most relevant snippets to prepend as system context.

支持视觉增强：自动识别图片文件，调用视觉 LLM 生成描述，将描述作为文本索引。
"""
from __future__ import annotations

import asyncio
import base64
import hashlib
import json
import math
import os
import re
from dataclasses import dataclass, field
from typing import Dict, List, Optional, Sequence, Tuple

from loguru import logger
from openai import OpenAI

try:  # Optional imports; code degrades gracefully if missing
    import docx  # type: ignore
except Exception:  # pragma: no cover - optional dependency
    docx = None

try:
    from pypdf import PdfReader  # type: ignore
except Exception:  # pragma: no cover - optional dependency
    PdfReader = None

try:
    import fitz  # PyMuPDF，用于提取 PDF 中的图片
except Exception:  # pragma: no cover - optional dependency
    fitz = None


@dataclass
class RAGConfig:
    doc_dir: str
    cache_path: str
    base_url: Optional[str]
    api_key: Optional[str]
    embedding_model: str = "text-embedding-3-small"
    # 独立的 embedding 配置，为空时回退到 base_url / api_key
    embedding_base_url: Optional[str] = None
    embedding_api_key: Optional[str] = None
    include_exts: Tuple[str, ...] = (".txt", ".md", ".log", ".docx", ".pdf")
    recursive: bool = True
    chunk_size: int = 900
    chunk_overlap: int = 150
    top_k: int = 4
    min_chunk_len: int = 120
    timeout: int = 60
    # embedding 专用配置
    embedding_timeout: int = 120
    embedding_batch_size: int = 16  # 每批发送的文本数
    embedding_max_retries: int = 2  # 失败重试次数
    embedding_max_tokens: int = 8000  # 单段文本最大 token 近似字符数
    
    # ========== 视觉增强配置 ==========
    enable_vision: bool = True  # 是否启用图片处理
    use_cached_images_only: bool = False  # 只使用已缓存的图片描述，不处理新图片
    vision_model: Optional[str] = None  # 视觉模型名称，为空时使用主 LLM（需在初始化时传入）
    vision_base_url: Optional[str] = None  # 视觉 API 地址，为空时使用 base_url
    vision_api_key: Optional[str] = None  # 视觉 API 密钥，为空时使用 api_key
    image_exts: Tuple[str, ...] = (".png", ".jpg", ".jpeg", ".gif", ".webp")  # 支持的图片格式
    vision_timeout: int = 120  # 视觉 API 超时时间
    vision_max_retries: int = 2  # 视觉 API 重试次数
    image_description_prompt: str = (
        "请分析这张图片，返回严格的 JSON 格式（不要添加 markdown 代码块）：\n"
        '{"is_relevant": true或false, "description": "描述内容"}\n\n'
        "判断 is_relevant=true 的标准（与电磁执行器/电机/线圈设计优化相关）：\n"
        "- 电磁结构、线圈、磁路、执行器的示意图或照片\n"
        "- 参数曲线、仿真结果、性能对比图表\n"
        "- 尺寸标注、公式推导、设计原理图\n"
        "- 磁场分布、力-位移曲线、效率图等\n\n"
        "判断 is_relevant=false 的标准（不相关）：\n"
        "- 封面、校名、徽标、签名页、声明页\n"
        "- 目录、参考文献列表、致谢\n"
        "- 纯文字截图、通用流程图、无关照片\n\n"
        "如果 is_relevant=true，description 需详细说明：尺寸参数、结构关系、数据趋势等。\n"
        "如果 is_relevant=false，description 简要说明图片内容即可。"
    )


# OpenAI 官方 embedding 端点
OPENAI_EMBEDDING_BASE_URL = "https://api.openai.com/v1"


class RAGEngine:
    def __init__(self, config: RAGConfig, vision_model: Optional[str] = None):
        """
        初始化 RAG 引擎。
        
        Args:
            config: RAG 配置
            vision_model: 视觉模型名称（如 gpt-5.1），为空时使用 config.vision_model
        """
        self.config = config
        # 为 embedding 单独创建 client，支持独立配置
        emb_base_url = config.embedding_base_url or config.base_url
        emb_api_key = config.embedding_api_key or config.api_key
        # 如果 base_url 不是 OpenAI 官方且未单独指定 embedding_base_url，
        # 则自动回退到 OpenAI 官方 embedding 端点（需要 OpenAI key）
        if emb_base_url and "openai.com" not in emb_base_url and not config.embedding_base_url:
            logger.info(
                f"检测到 LLM base_url={emb_base_url} 非 OpenAI，embedding 自动切换到 {OPENAI_EMBEDDING_BASE_URL}"
            )
            emb_base_url = OPENAI_EMBEDDING_BASE_URL
            # 尝试从环境变量获取 OpenAI key
            emb_api_key = config.embedding_api_key or os.getenv("OPENAI_API_KEY") or config.api_key
        self._emb_base_url = emb_base_url
        self._emb_api_key = emb_api_key
        self.embedding_client = OpenAI(
            api_key=emb_api_key,
            base_url=emb_base_url,
            timeout=config.embedding_timeout,
        )
        self.index: List[Dict] = []
        self._prepared = False
        
        # ========== 视觉增强：初始化视觉 client ==========
        self._vision_model = vision_model or config.vision_model
        self._vision_client: Optional[OpenAI] = None
        self._image_cache: Dict[str, str] = {}  # {image_hash: description}
        # 从 cache_path 推导图片缓存路径
        cache_dir = os.path.dirname(config.cache_path) if config.cache_path else "cache"
        self._image_cache_path = os.path.join(cache_dir, "image_descriptions_cache.json")
        
        # 尝试加载已有的图片描述缓存
        self._load_image_cache()
        
        if config.enable_vision and self._vision_model:
            vision_base_url = config.vision_base_url or config.base_url
            vision_api_key = config.vision_api_key or config.api_key
            self._vision_client = OpenAI(
                api_key=vision_api_key,
                base_url=vision_base_url,
                timeout=config.vision_timeout,
            )
            logger.info(
                f"RAGEngine 视觉增强已启用 | vision_model={self._vision_model} | base_url={vision_base_url}"
            )
        
        logger.info(
            f"RAGEngine 初始化完成 | embedding_base_url={emb_base_url} | model={config.embedding_model}"
        )

    async def prepare(self) -> None:
        """Load cached embeddings or build a fresh index on a background thread."""
        await asyncio.to_thread(self._ensure_index)

    async def build_context(self, query: str, top_k: Optional[int] = None) -> str:
        """Return formatted context for a query; empty string if unavailable."""
        await self.prepare()
        if not self.index:
            return ""
        try:
            snippets = await asyncio.to_thread(self._retrieve, query, top_k)
        except Exception as exc:  # pragma: no cover - runtime guard
            logger.error(f"RAG 检索失败: {exc}")
            return ""
        if not snippets:
            return ""

        lines = [
            "以下是从本地文献中检索到的相关片段，请基于这些内容回答，不要编造文献未提及的信息："
        ]
        for item in snippets:
            trimmed = item["text"].strip()
            lines.append(f"[来源: {item['source']} | 相似度: {item['score']:.3f}]\n{trimmed}")
        return "\n\n".join(lines)

    # --- 图片描述解析与缓存管理 -------------------------------------------------
    
    def _parse_image_response(self, raw_response: str) -> Tuple[bool, str]:
        """
        解析 LLM 返回的图片描述 JSON。
        
        Returns:
            (is_relevant, description)
        """
        import re
        
        # 尝试直接解析 JSON
        try:
            # 移除可能的 markdown 代码块标记
            cleaned = raw_response.strip()
            if cleaned.startswith("```"):
                # 移除 ```json 或 ``` 开头和结尾
                cleaned = re.sub(r'^```\w*\n?', '', cleaned)
                cleaned = re.sub(r'\n?```$', '', cleaned)
            
            data = json.loads(cleaned)
            is_relevant = data.get("is_relevant", True)
            description = data.get("description", "")
            return (is_relevant, description)
        except json.JSONDecodeError:
            pass
        
        # 如果 JSON 解析失败，尝试从文本中提取
        try:
            # 尝试找到 JSON 部分
            match = re.search(r'\{[^{}]*"is_relevant"\s*:\s*(true|false)[^{}]*\}', raw_response, re.IGNORECASE)
            if match:
                data = json.loads(match.group(0))
                return (data.get("is_relevant", True), data.get("description", raw_response))
        except Exception:
            pass
        
        # 最后回退：假设相关，使用原始响应作为描述
        logger.warning(f"无法解析图片描述 JSON，假设相关: {raw_response[:100]}...")
        return (True, raw_response)
    
    def _load_image_cache(self) -> None:
        """启动时加载已有的图片描述缓存"""
        if os.path.exists(self._image_cache_path):
            try:
                with open(self._image_cache_path, "r", encoding="utf-8") as f:
                    self._image_cache = json.load(f)
                logger.info(f"📂 加载图片描述缓存: {len(self._image_cache)} 条 | {self._image_cache_path}")
            except Exception as e:
                logger.warning(f"加载图片描述缓存失败: {e}")
                self._image_cache = {}
    
    def _save_image_cache(self) -> None:
        """保存图片描述缓存（每次新增描述后调用）"""
        try:
            os.makedirs(os.path.dirname(self._image_cache_path), exist_ok=True)
            with open(self._image_cache_path, "w", encoding="utf-8") as f:
                json.dump(self._image_cache, f, ensure_ascii=False, indent=2)
            logger.debug(f"💾 图片描述缓存已保存: {len(self._image_cache)} 条")
        except Exception as e:
            logger.warning(f"保存图片描述缓存失败: {e}")

    # --- internal helpers -------------------------------------------------
    def _ensure_index(self) -> None:
        if self._prepared:
            return
        cache_path = self.config.cache_path
        if cache_path and os.path.exists(cache_path):
            try:
                with open(cache_path, "r", encoding="utf-8") as f:
                    payload = json.load(f)
                self.index = payload.get("index", [])
                # 恢复图片描述缓存
                self._image_cache = payload.get("image_cache", {})
                image_count = len(self._image_cache)
                logger.info(f"RAG 从缓存加载 {len(self.index)} 个片段 (含 {image_count} 张图片描述): {cache_path}")
                self._prepared = True
                return
            except Exception as exc:  # pragma: no cover - cache may be stale
                logger.warning(f"RAG 缓存读取失败，将重新构建：{exc}")

        documents = self._load_documents(self.config.doc_dir)
        if not documents:
            logger.warning("RAG 未在文献目录中找到可用文件")
            self.index = []
            self._prepared = True
            return

        chunks = self._chunk_documents(documents)
        if not chunks:
            logger.warning("RAG 文档分块为空")
            self.index = []
            self._prepared = True
            return

        embeddings = self._embed_texts([c["text"] for c in chunks])
        failed = sum(1 for emb in embeddings if not emb)
        if failed == len(embeddings):
            logger.warning("Embedding 全部失败，将仅使用词法匹配回退")
        elif failed:
            logger.warning(f"部分 Embedding 失败：{failed}/{len(embeddings)}，向量检索将回退词法匹配")

        self.index = [
            {"text": chunk["text"], "source": chunk["source"], "embedding": emb}
            for chunk, emb in zip(chunks, embeddings)
        ]
        self._prepared = True

        if cache_path:
            try:
                with open(cache_path, "w", encoding="utf-8") as f:
                    # 同时保存图片描述缓存
                    json.dump({
                        "model": self.config.embedding_model,
                        "index": self.index,
                        "image_cache": self._image_cache,
                    }, f, ensure_ascii=False)
                logger.info(f"RAG 缓存写入完成: {cache_path}")
            except Exception as exc:  # pragma: no cover - cache writing best-effort
                logger.warning(f"RAG 缓存写入失败：{exc}")

    def _load_documents(self, doc_dir: str) -> List[Tuple[str, str]]:
        if not doc_dir:
            return []

        # 支持通过分号/竖线分隔的多目录配置，方便同时索引 liter 与日志目录
        raw_dirs = re.split(r"[;|]", doc_dir)
        doc_roots = [d.strip() for d in raw_dirs if d.strip() and os.path.isdir(d.strip())]
        if not doc_roots:
            return []

        docs: List[Tuple[str, str]] = []
        embedded_images: List[Tuple[str, str]] = []  # 从文档中提取的嵌入图片
        exts = {ext.lower() for ext in self.config.include_exts}
        
        # 判断是否只使用缓存的图片
        use_cached_only = self.config.use_cached_images_only
        if use_cached_only:
            logger.info("📦 使用缓存模式：跳过新图片处理，只使用已缓存的图片描述")
        
        for root in doc_roots:
            walker = os.walk(root) if self.config.recursive else ((root, [], os.listdir(root)),)
            for dirpath, _, filenames in walker:
                for file_name in filenames:
                    full_path = os.path.join(dirpath, file_name)
                    if not os.path.isfile(full_path):
                        continue
                    ext = os.path.splitext(file_name)[1].lower()
                    if ext not in exts:
                        continue

                    rel_name = os.path.relpath(full_path, root)
                    if ext in {".txt", ".md", ".log"}:
                        docs.append((rel_name, self._read_text_file(full_path)))
                    elif ext == ".docx":
                        docs.append((rel_name, self._read_docx(full_path)))
                        # ★ 视觉增强：同时提取 DOCX 中的嵌入图片（除非只用缓存）
                        if self.config.enable_vision and not use_cached_only:
                            embedded_images.extend(self._extract_images_from_docx(full_path, rel_name))
                    elif ext == ".pdf":
                        docs.append((rel_name, self._read_pdf(full_path)))
                        # ★ 视觉增强：同时提取 PDF 中的嵌入图片（除非只用缓存）
                        if self.config.enable_vision and not use_cached_only:
                            embedded_images.extend(self._extract_images_from_pdf(full_path, rel_name))
        
        # ========== 视觉增强：加载独立图片文件 ==========
        standalone_images: List[Tuple[str, str]] = []
        if self.config.enable_vision and not use_cached_only:
            standalone_images = self._load_images(doc_roots)
        
        # ========== 使用缓存的图片描述 ==========
        cached_images: List[Tuple[str, str]] = []
        if self.config.enable_vision and use_cached_only:
            # 从缓存中加载所有非 SKIP 的图片描述
            for img_hash, description in self._image_cache.items():
                if not description.startswith("[SKIP]") and description.strip():
                    cached_images.append((f"cached_img_{img_hash[:8]}", description))
            logger.info(f"📦 从缓存加载了 {len(cached_images)} 张图片描述")
        
        # 合并所有图片描述
        all_images = embedded_images + standalone_images + cached_images
        if all_images:
            docs.extend(all_images)
            if use_cached_only:
                logger.info(f"RAG 视觉增强: 使用缓存图片 {len(cached_images)} 张")
            else:
                logger.info(f"RAG 视觉增强: 嵌入图片 {len(embedded_images)} 张 + 独立图片 {len(standalone_images)} 张")

        return [(name, text) for name, text in docs if text.strip()]

    def _read_text_file(self, path: str) -> str:
        try:
            with open(path, "r", encoding="utf-8") as f:
                return f.read()
        except Exception:
            try:
                with open(path, "r", encoding="gbk", errors="ignore") as f:
                    return f.read()
            except Exception as exc:  # pragma: no cover - best effort
                logger.warning(f"读取文本失败: {path} | {exc}")
                return ""

    def _read_docx(self, path: str) -> str:
        if docx is None:
            logger.warning("未安装 python-docx，无法读取 DOCX，运行 `pip install python-docx` 启用")
            return ""
        try:
            document = docx.Document(path)
            return "\n".join(p.text for p in document.paragraphs)
        except Exception as exc:  # pragma: no cover - parsing edge cases
            logger.warning(f"解析 DOCX 失败: {path} | {exc}")
            return ""

    def _read_pdf(self, path: str) -> str:
        if PdfReader is None:
            logger.warning("未安装 pypdf，无法读取 PDF，运行 `pip install pypdf` 启用")
            return ""
        try:
            reader = PdfReader(path)
            pages = [page.extract_text() or "" for page in reader.pages]
            return "\n".join(pages)
        except Exception as exc:  # pragma: no cover - parsing edge cases
            logger.warning(f"解析 PDF 失败: {path} | {exc}")
            return ""

    # ========== 视觉增强：图片处理方法 ==========
    
    def _get_image_hash(self, path: str) -> str:
        """计算图片文件的 MD5 哈希，用于缓存标识"""
        try:
            with open(path, "rb") as f:
                return hashlib.md5(f.read()).hexdigest()
        except Exception:
            return ""
    
    def _read_image_as_base64(self, path: str) -> Tuple[Optional[str], str]:
        """
        读取图片并转换为 base64。
        
        Returns:
            (base64_data, media_type) 或 (None, "") 如果失败
        """
        ext = os.path.splitext(path)[1].lower()
        media_type_map = {
            ".png": "image/png",
            ".jpg": "image/jpeg",
            ".jpeg": "image/jpeg",
            ".gif": "image/gif",
            ".webp": "image/webp",
        }
        media_type = media_type_map.get(ext, "image/png")
        
        try:
            with open(path, "rb") as f:
                data = base64.b64encode(f.read()).decode("utf-8")
            return data, media_type
        except Exception as exc:
            logger.warning(f"读取图片失败: {path} | {exc}")
            return None, ""
    
    def _describe_image(self, path: str, rel_name: str) -> str:
        """
        调用视觉 LLM 生成图片描述。
        
        Args:
            path: 图片完整路径
            rel_name: 相对路径（用于日志和缓存标识）
        
        Returns:
            图片描述文本，失败返回空字符串
        """
        if not self._vision_client or not self._vision_model:
            return ""
        
        # 检查缓存
        img_hash = self._get_image_hash(path)
        if img_hash and img_hash in self._image_cache:
            cached = self._image_cache[img_hash]
            if cached.startswith("[SKIP]"):
                # 之前判断为不相关，跳过
                logger.debug(f"⏭️ 图片缓存命中(跳过): {rel_name}")
                return ""
            logger.debug(f"图片描述命中缓存: {rel_name}")
            return cached
        
        # 读取图片
        base64_data, media_type = self._read_image_as_base64(path)
        if not base64_data:
            return ""
        
        # 调用视觉 API
        import time
        max_retries = self.config.vision_max_retries
        
        for attempt in range(max_retries + 1):
            try:
                response = self._vision_client.chat.completions.create(
                    model=self._vision_model,
                    messages=[
                        {
                            "role": "user",
                            "content": [
                                {
                                    "type": "text",
                                    "text": f"图片文件名: {rel_name}\n\n{self.config.image_description_prompt}"
                                },
                                {
                                    "type": "image_url",
                                    "image_url": {
                                        "url": f"data:{media_type};base64,{base64_data}"
                                    }
                                }
                            ]
                        }
                    ],
                    max_tokens=1000,
                )
                
                raw_response = response.choices[0].message.content or ""
                if raw_response:
                    # 解析 JSON 响应，判断相关性
                    is_relevant, description = self._parse_image_response(raw_response)
                    
                    if not is_relevant:
                        # 不相关的图片，记录但不索引
                        logger.info(f"⏭️ 跳过不相关图片: {rel_name}")
                        # 仍然缓存（避免重复调用 API），但标记为跳过
                        if img_hash:
                            self._image_cache[img_hash] = f"[SKIP] {description}"
                            self._save_image_cache()
                        return ""  # 返回空，不加入索引
                    
                    # 相关图片，添加来源标识
                    description = f"[图片: {rel_name}]\n{description}"
                    # 缓存结果并立即保存
                    if img_hash:
                        self._image_cache[img_hash] = description
                        self._save_image_cache()
                    logger.info(f"✅ 图片描述生成成功: {rel_name} ({len(description)} 字符)")
                    return description
                else:
                    logger.warning(f"图片描述为空: {rel_name}")
                    return ""
                    
            except Exception as exc:
                if attempt < max_retries:
                    wait = 2 ** attempt
                    logger.warning(
                        f"图片描述生成失败 (尝试 {attempt + 1}/{max_retries + 1})，{wait}s 后重试: {rel_name} | {exc}"
                    )
                    time.sleep(wait)
                else:
                    logger.error(f"图片描述生成最终失败: {rel_name} | {exc}")
        
        return ""
    
    def _load_images(self, doc_roots: List[str]) -> List[Tuple[str, str]]:
        """
        扫描目录中的图片文件并生成描述。
        
        Returns:
            [(rel_name, description), ...]
        """
        if not self.config.enable_vision or not self._vision_client:
            return []
        
        image_exts = {ext.lower() for ext in self.config.image_exts}
        images: List[Tuple[str, str]] = []
        
        for root in doc_roots:
            walker = os.walk(root) if self.config.recursive else ((root, [], os.listdir(root)),)
            for dirpath, _, filenames in walker:
                for file_name in filenames:
                    full_path = os.path.join(dirpath, file_name)
                    if not os.path.isfile(full_path):
                        continue
                    ext = os.path.splitext(file_name)[1].lower()
                    if ext not in image_exts:
                        continue
                    
                    rel_name = os.path.relpath(full_path, root)
                    logger.info(f"🖼️ 处理图片: {rel_name}")
                    description = self._describe_image(full_path, rel_name)
                    if description:
                        images.append((rel_name, description))
        
        if images:
            logger.info(f"视觉增强完成: 成功处理 {len(images)} 张独立图片")
        
        return images
    
    # ========== 视觉增强：从文档中提取嵌入图片 ==========
    
    def _extract_images_from_pdf(self, path: str, rel_name: str) -> List[Tuple[str, str]]:
        """
        从 PDF 文件中提取嵌入的图片并生成描述。
        
        Returns:
            [(image_name, description), ...]
        """
        if fitz is None:
            logger.debug("未安装 pymupdf，无法提取 PDF 中的图片，运行 `pip install pymupdf` 启用")
            return []
        
        if not self._vision_client:
            return []
        
        results: List[Tuple[str, str]] = []
        
        try:
            doc = fitz.open(path)
            image_count = 0
            
            for page_num, page in enumerate(doc):
                image_list = page.get_images(full=True)
                
                for img_index, img_info in enumerate(image_list):
                    xref = img_info[0]
                    
                    try:
                        # 提取图片数据
                        base_image = doc.extract_image(xref)
                        image_bytes = base_image["image"]
                        image_ext = base_image.get("ext", "png")
                        
                        # 跳过太小的图片（可能是图标或装饰）
                        if len(image_bytes) < 5000:  # 小于 5KB
                            continue
                        
                        # 生成图片标识
                        image_name = f"{rel_name}#page{page_num + 1}_img{img_index + 1}"
                        
                        # 检查缓存
                        img_hash = hashlib.md5(image_bytes).hexdigest()
                        if img_hash in self._image_cache:
                            cached = self._image_cache[img_hash]
                            if cached.startswith("[SKIP]"):
                                # 之前判断为不相关，跳过
                                logger.debug(f"⏭️ PDF 图片缓存命中(跳过): {image_name}")
                                continue
                            logger.debug(f"PDF 图片描述命中缓存: {image_name}")
                            results.append((image_name, cached))
                            image_count += 1
                            continue
                        
                        # 转 base64
                        base64_data = base64.b64encode(image_bytes).decode("utf-8")
                        media_type = f"image/{image_ext}" if image_ext in ["png", "jpeg", "jpg", "gif", "webp"] else "image/png"
                        
                        # 调用视觉 API（传入 img_hash 用于缓存跳过的图片）
                        description = self._describe_image_from_base64(
                            base64_data, media_type, image_name, img_hash
                        )
                        
                        if description:
                            # 相关图片，保存并加入结果
                            self._image_cache[img_hash] = description
                            self._save_image_cache()
                            results.append((image_name, description))
                            image_count += 1
                            # 注意：不相关的图片已在 _describe_image_from_base64 中缓存为 [SKIP]
                            
                    except Exception as e:
                        logger.debug(f"提取 PDF 图片失败: {image_name} | {e}")
                        continue
            
            doc.close()
            
            if image_count > 0:
                logger.info(f"📄 从 PDF 提取了 {image_count} 张图片: {rel_name}")
                
        except Exception as exc:
            logger.warning(f"读取 PDF 失败: {path} | {exc}")
        
        return results
    
    def _extract_images_from_docx(self, path: str, rel_name: str) -> List[Tuple[str, str]]:
        """
        从 DOCX 文件中提取嵌入的图片并生成描述。
        
        Returns:
            [(image_name, description), ...]
        """
        if docx is None:
            logger.debug("未安装 python-docx，无法提取 DOCX 中的图片")
            return []
        
        if not self._vision_client:
            return []
        
        results: List[Tuple[str, str]] = []
        
        try:
            document = docx.Document(path)
            image_count = 0
            
            # 遍历文档中的所有关系，找到图片
            for rel_id, rel in document.part.rels.items():
                if "image" in rel.reltype:
                    try:
                        image_part = rel.target_part
                        image_bytes = image_part.blob
                        
                        # 跳过太小的图片
                        if len(image_bytes) < 5000:
                            continue
                        
                        # 获取图片格式
                        content_type = image_part.content_type
                        if "png" in content_type:
                            media_type = "image/png"
                        elif "jpeg" in content_type or "jpg" in content_type:
                            media_type = "image/jpeg"
                        elif "gif" in content_type:
                            media_type = "image/gif"
                        elif "webp" in content_type:
                            media_type = "image/webp"
                        else:
                            media_type = "image/png"
                        
                        # 生成图片标识
                        image_count += 1
                        image_name = f"{rel_name}#img{image_count}"
                        
                        # 检查缓存
                        img_hash = hashlib.md5(image_bytes).hexdigest()
                        if img_hash in self._image_cache:
                            cached = self._image_cache[img_hash]
                            if cached.startswith("[SKIP]"):
                                # 之前判断为不相关，跳过
                                logger.debug(f"⏭️ DOCX 图片缓存命中(跳过): {image_name}")
                                continue
                            logger.debug(f"DOCX 图片描述命中缓存: {image_name}")
                            results.append((image_name, cached))
                            continue
                        
                        # 转 base64
                        base64_data = base64.b64encode(image_bytes).decode("utf-8")
                        
                        # 调用视觉 API（传入 img_hash 用于缓存跳过的图片）
                        description = self._describe_image_from_base64(
                            base64_data, media_type, image_name, img_hash
                        )
                        
                        if description:
                            # 相关图片，保存并加入结果
                            self._image_cache[img_hash] = description
                            self._save_image_cache()
                            results.append((image_name, description))
                            # 注意：不相关的图片已在 _describe_image_from_base64 中缓存为 [SKIP]
                            
                    except Exception as e:
                        logger.debug(f"提取 DOCX 图片失败: {rel_id} | {e}")
                        continue
            
            if image_count > 0:
                logger.info(f"📝 从 DOCX 提取了 {len(results)} 张图片: {rel_name}")
                
        except Exception as exc:
            logger.warning(f"读取 DOCX 失败: {path} | {exc}")
        
        return results
    
    def _describe_image_from_base64(
        self, base64_data: str, media_type: str, image_name: str, img_hash: Optional[str] = None
    ) -> str:
        """
        从 base64 数据调用视觉 LLM 生成描述。
        
        Returns:
            描述文本（相关图片）或空字符串（不相关或失败）
        """
        if not self._vision_client or not self._vision_model:
            return ""
        
        import time
        max_retries = self.config.vision_max_retries
        
        for attempt in range(max_retries + 1):
            try:
                response = self._vision_client.chat.completions.create(
                    model=self._vision_model,
                    messages=[
                        {
                            "role": "user",
                            "content": [
                                {
                                    "type": "text",
                                    "text": f"图片来源: {image_name}\n\n{self.config.image_description_prompt}"
                                },
                                {
                                    "type": "image_url",
                                    "image_url": {
                                        "url": f"data:{media_type};base64,{base64_data}"
                                    }
                                }
                            ]
                        }
                    ],
                    max_tokens=1000,
                )
                
                raw_response = response.choices[0].message.content or ""
                if raw_response:
                    # 解析 JSON 响应，判断相关性
                    is_relevant, description = self._parse_image_response(raw_response)
                    
                    if not is_relevant:
                        # 不相关的图片，记录但不索引
                        logger.info(f"⏭️ 跳过不相关图片: {image_name}")
                        # 缓存标记为跳过（避免重复调用 API）
                        if img_hash:
                            self._image_cache[img_hash] = f"[SKIP] {description}"
                            self._save_image_cache()
                        return ""  # 返回空，不加入索引
                    
                    # 相关图片
                    description = f"[图片: {image_name}]\n{description}"
                    logger.info(f"✅ 图片描述生成成功: {image_name}")
                    return description
                    
            except Exception as exc:
                if attempt < max_retries:
                    wait = 2 ** attempt
                    logger.warning(
                        f"图片描述生成失败 (尝试 {attempt + 1}/{max_retries + 1})，{wait}s 后重试: {image_name} | {exc}"
                    )
                    time.sleep(wait)
                else:
                    logger.error(f"图片描述生成最终失败: {image_name} | {exc}")
        
        return ""

    def _chunk_documents(self, documents: Sequence[Tuple[str, str]]) -> List[Dict]:
        chunks: List[Dict] = []
        size = max(100, self.config.chunk_size)
        overlap = max(0, min(self.config.chunk_overlap, size - 1))
        min_len = max(40, self.config.min_chunk_len)
        for name, text in documents:
            cleaned = self._clean_text(text)
            if not cleaned:
                continue
            start = 0
            while start < len(cleaned):
                end = start + size
                piece = cleaned[start:end]
                if len(piece) >= min_len:
                    chunks.append({"text": piece, "source": name})
                if end >= len(cleaned):
                    break
                start = end - overlap
        return chunks

    def _embed_texts(self, texts: Sequence[str]) -> List[Optional[List[float]]]:
        """生成文本嵌入向量，支持分批、重试、截断。"""
        if not texts:
            return []

        # 预处理：截断过长文本，跳过空文本
        max_chars = self.config.embedding_max_tokens
        processed: List[Tuple[int, str]] = []  # (原始索引, 处理后文本)
        for i, text in enumerate(texts):
            t = text.strip()
            if not t:
                continue
            if len(t) > max_chars:
                t = t[:max_chars]
            processed.append((i, t))

        if not processed:
            return [None for _ in texts]

        result: List[Optional[List[float]]] = [None] * len(texts)
        batch_size = max(1, self.config.embedding_batch_size)
        max_retries = max(0, self.config.embedding_max_retries)
        total_success = 0
        total_fail = 0

        # 分批处理
        for batch_start in range(0, len(processed), batch_size):
            batch = processed[batch_start : batch_start + batch_size]
            batch_texts = [t for _, t in batch]
            batch_indices = [i for i, _ in batch]

            embeddings = self._embed_batch_with_retry(batch_texts, max_retries)
            for idx, emb in zip(batch_indices, embeddings):
                result[idx] = emb
                if emb:
                    total_success += 1
                else:
                    total_fail += 1

        logger.info(
            f"Embedding 完成 | 成功={total_success} 失败={total_fail} 总计={len(texts)} "
            f"| base_url={self._emb_base_url} model={self.config.embedding_model}"
        )
        return result

    def _embed_batch_with_retry(
        self, texts: List[str], max_retries: int
    ) -> List[Optional[List[float]]]:
        """对一批文本调用 embedding API，支持重试。"""
        import time

        last_exc: Optional[Exception] = None
        for attempt in range(max_retries + 1):
            try:
                response = self.embedding_client.embeddings.create(
                    model=self.config.embedding_model,
                    input=texts,
                )
                # 校验响应
                if not response or not hasattr(response, "data") or not response.data:
                    raise ValueError("No embedding data received from API")

                # 按 index 排序，确保顺序正确
                sorted_data = sorted(response.data, key=lambda x: getattr(x, "index", 0))
                embeddings: List[Optional[List[float]]] = []
                for item in sorted_data:
                    emb = getattr(item, "embedding", None)
                    if emb and isinstance(emb, list) and len(emb) > 0:
                        embeddings.append(emb)
                    else:
                        embeddings.append(None)

                # 长度校验
                if len(embeddings) != len(texts):
                    logger.warning(
                        f"Embedding 返回数量不匹配：期望 {len(texts)}，实际 {len(embeddings)}"
                    )
                    # 补齐 None
                    while len(embeddings) < len(texts):
                        embeddings.append(None)

                return embeddings

            except Exception as exc:
                last_exc = exc
                if attempt < max_retries:
                    wait = 2 ** attempt  # 指数退避：1s, 2s, 4s...
                    logger.warning(
                        f"Embedding 请求失败 (尝试 {attempt + 1}/{max_retries + 1})，{wait}s 后重试: {exc}"
                    )
                    time.sleep(wait)
                else:
                    # 最后一次失败，记录详细错误
                    error_msg = str(exc)
                    # 尝试提取更多信息
                    if hasattr(exc, "response"):
                        try:
                            resp = exc.response
                            status = getattr(resp, "status_code", "unknown")
                            body = getattr(resp, "text", "")[:500]
                            error_msg = f"status={status} body={body}"
                        except Exception:
                            pass
                    logger.error(
                        f"Embedding 请求最终失败 | base_url={self._emb_base_url} "
                        f"model={self.config.embedding_model} | {error_msg}"
                    )

        return [None for _ in texts]

    def _retrieve(self, query: str, top_k: Optional[int]) -> List[Dict]:
        if not self.index:
            return []
        query_emb = self._embed_texts([query])[0]
        embedding_items = [item for item in self.index if item.get("embedding")]

        if query_emb and embedding_items:
            scored = [
                {
                    "text": item["text"],
                    "source": item["source"],
                    "score": self._cosine_similarity(query_emb, item["embedding"]),
                }
                for item in embedding_items
            ]
        else:
            scored = self._lexical_score(query)
        scored.sort(key=lambda x: x["score"], reverse=True)
        k = top_k or self.config.top_k
        return scored[:max(1, k)]

    def _lexical_score(self, query: str) -> List[Dict]:
        tokens = set(self._tokenize(query))
        scored: List[Dict] = []
        if not tokens:
            return scored
        for item in self.index:
            doc_tokens = set(self._tokenize(item["text"]))
            inter = len(tokens & doc_tokens)
            score = inter / (len(tokens) + 1e-6)
            scored.append({"text": item["text"], "source": item["source"], "score": score})
        return scored

    def _cosine_similarity(self, a: Sequence[float], b: Sequence[float]) -> float:
        if not a or not b or len(a) != len(b):
            return 0.0
        dot = sum(x * y for x, y in zip(a, b))
        norm_a = math.sqrt(sum(x * x for x in a))
        norm_b = math.sqrt(sum(y * y for y in b))
        if norm_a == 0 or norm_b == 0:
            return 0.0
        return dot / (norm_a * norm_b)

    def _clean_text(self, text: str) -> str:
        text = re.sub(r"\s+", " ", text)
        return text.strip()

    def _tokenize(self, text: str) -> List[str]:
        return [t for t in re.split(r"[^\w]+", text.lower()) if t]
